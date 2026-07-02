from __future__ import annotations

import argparse
import json
import logging
import re
import shutil
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt

from cmie.core.unit_config import UnitConfig
from cmie.generator import assessment_generator
from cmie.generator.ai_lesson_engine import sanitize_ai_text, slugify
from cmie.generator.assessment_markdown import render_assessment_markdown
from cmie.generator.batch_generate import run_batch
from cmie.generator.canva_csv_generator import export_unit_canva_csv
from cmie.generator.readme_generator import generate_unit_readme
from cmie.generator.roadmap_generator import generate_unit_roadmap
from cmie.generator.workbook_generator import generate_student_workbook
from cmie.marketing.marketing_generator import generate_marketing_assets
from cmie.validation.unit_validation import validate_unit

try:
    from cmie.generator.listing_generator import generate_listings_for_unit  # type: ignore
except Exception:  # pragma: no cover
    generate_listings_for_unit = None  # type: ignore


# --------------------------------------------------------------------
# Logging
# --------------------------------------------------------------------


def setup_logger() -> logging.Logger:
    logger = logging.getLogger("full_product_pipeline")
    if not logger.handlers:
        logger.setLevel(logging.INFO)
        handler = logging.StreamHandler()
        formatter = logging.Formatter("[%(levelname)s] %(message)s")
        handler.setFormatter(formatter)
        logger.addHandler(handler)
    return logger


# --------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------


def markdown_to_docx(md_path: Path, docx_path: Path, logger: logging.Logger) -> None:
    """
    Convert simple markdown-like workbook/task files into a cleaner DOCX.
    Supports:
    - #, ##, ### headings
    - bullet points
    - horizontal rules
    - response box blocks triggered by [Write your response below]
    - page breaks before each new lesson
    """
    logger.info(f"  Converting {md_path.name} -> {docx_path.name}")
    text = md_path.read_text(encoding="utf-8")
    # Defensive: strip stray control characters (vertical tabs etc.) that
    # occasionally survive in AI-generated markdown, so they never reach
    # the rendered DOCX.
    text = sanitize_ai_text(text)
    lines = text.splitlines()

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(11)

    def add_response_box(rows: int = 5) -> None:
        table = doc.add_table(rows=rows, cols=1)
        table.style = "Table Grid"

        for row in table.rows:
            row.height = Inches(0.6)
            cell = row.cells[0]
            cell.text = ""

    inline_md_re = re.compile(
        r"(\*\*.+?\*\*"          # **bold**
        r"|`[^`\n]+`"            # `inline code`
        r"|(?<!\*)\*[^*\n]+\*(?!\*))"  # *italic* (not part of **)
    )
    link_re = re.compile(r"\[([^\]]+)\]\(([^)\s]+)\)")

    def add_runs_with_bold(paragraph, text: str) -> None:
        """Convert inline markdown into real DOCX runs instead of leaving
        literal markers in the output. Handles **bold**, *italic*,
        `inline code` (monospace run), and [text](url) links (rendered as
        'text (url)'). Previously only block-level markdown was handled,
        so these markers rendered literally in every generated docx
        (roadmap, assessment, teacher guide, workbook)."""
        text = link_re.sub(r"\1 (\2)", text)
        for part in inline_md_re.split(text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**") and len(part) > 4:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("`") and part.endswith("`") and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Consolas"
            elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

    def is_table_row(line: str) -> bool:
        return line.strip().startswith("|") and line.strip().endswith("|")

    def is_table_separator(line: str) -> bool:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        return all(c == "" or set(c) <= {"-", ":"} for c in cells) and len(cells) > 0

    def add_markdown_table(rows: List[str]) -> None:
        parsed = [
            [c.strip() for c in row.strip().strip("|").split("|")]
            for row in rows
            if not is_table_separator(row)
        ]
        if not parsed:
            return
        n_cols = max(len(r) for r in parsed)
        table = doc.add_table(rows=len(parsed), cols=n_cols)
        table.style = "Table Grid"
        for r, row_cells in enumerate(parsed):
            for c in range(n_cols):
                cell_text = row_cells[c] if c < len(row_cells) else ""
                cell = table.cell(r, c)
                cell.text = ""
                add_runs_with_bold(cell.paragraphs[0], cell_text)
                if r == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

    def section_is_empty(heading_idx: int) -> bool:
        """True when a heading has no content before the next heading,
        horizontal rule, or end of file. Used to skip rendering headings
        like 'Extension ideas:' that AI output sometimes leaves without
        any body."""
        j = heading_idx + 1
        while j < len(lines) and not lines[j].strip():
            j += 1
        if j >= len(lines):
            return True
        nxt = lines[j].strip()
        return nxt.startswith("#") or nxt == "---"

    previous_was_blank = True
    i = 0
    while i < len(lines):
        raw_line = lines[i]
        stripped = raw_line.strip()

        if not stripped:
            if not previous_was_blank:
                doc.add_paragraph("")
            previous_was_blank = True
            i += 1
            continue

        previous_was_blank = False

        if is_table_row(stripped):
            table_lines = []
            while i < len(lines) and is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(table_lines)
            doc.add_paragraph("")
            continue

        if stripped == "---":
            doc.add_paragraph("")
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Title")
            add_runs_with_bold(p, stripped[2:].strip())
            i += 1
            continue

        if stripped.startswith("## "):
            heading_text = stripped[3:].strip()

            if heading_text.upper().startswith("LESSON "):
                if len(doc.paragraphs) > 0:
                    doc.add_page_break()

            p = doc.add_paragraph(style="Heading 1")
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(16)
            i += 1
            continue

        if stripped.startswith("### "):
            if section_is_empty(i):
                i += 1
                continue
            p = doc.add_paragraph(style="Heading 2")
            add_runs_with_bold(p, stripped[4:].strip())
            doc.add_paragraph("")
            i += 1
            continue

        if stripped.startswith("#### "):
            # Previously unhandled: '#### ' headings fell through and were
            # rendered as literal '#### text' paragraphs.
            if section_is_empty(i):
                i += 1
                continue
            p = doc.add_paragraph(style="Heading 3")
            add_runs_with_bold(p, stripped[5:].strip())
            i += 1
            continue

        if stripped == "[Write your response below]":
            p = doc.add_paragraph()
            run = p.add_run("Write your response below:")
            run.italic = True
            add_response_box(rows=5)
            doc.add_paragraph("")
            i += 1
            continue

        if set(stripped) == {"_"}:
            i += 1
            continue

        bullet_match = re.match(r"[-*+]\s+", stripped)
        if bullet_match:
            # AI-generated markdown sometimes flattens a nested bullet into
            # "- - text" instead of proper indentation -- strip repeated
            # leading bullet markers rather than leaving a literal "- " in
            # the rendered text. Also accept "*" / "+" bullet markers and
            # preserve one level of nesting via indentation.
            bullet_text = stripped[bullet_match.end():].strip()
            while bullet_text.startswith("- "):
                bullet_text = bullet_text[2:].strip()
            indent = len(raw_line.expandtabs(4)) - len(raw_line.expandtabs(4).lstrip())
            style_name = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style_name)
            add_runs_with_bold(p, bullet_text)
            i += 1
            continue

        # NOTE: numbered lists ("1. step") are deliberately left as plain
        # paragraphs with their literal marker -- python-docx's "List
        # Number" style shares one numbering sequence document-wide, so
        # separate lists would run 1..4 then 5..7. The literal marker
        # reads correctly and is the safer rendering.

        p = doc.add_paragraph()
        add_runs_with_bold(p, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


# --------------------------------------------------------------------
# Stages
# --------------------------------------------------------------------


def stage_lessons(
    cfg: UnitConfig,
    unit_root: Path,
    logger: logging.Logger,
) -> Tuple[Path, List[Path]]:
    """
    Generate lessons and copy lesson JSON files into the unit release folder.
    Slides are now handled via CSV.
    """
    logger.info("Stage: lessons")

    lessons_dir = unit_root / "lessons"
    lessons_dir.mkdir(parents=True, exist_ok=True)

    for p in lessons_dir.glob("*.json"):
        p.unlink()

    results = run_batch(cfg)

    copied_lessons: List[Path] = []
    for lesson_json in results:
        topic_slug = lesson_json.parent.name
        dest_json = lessons_dir / f"{topic_slug}.json"
        shutil.copy2(lesson_json, dest_json)
        copied_lessons.append(dest_json)

    logger.info(f"  {len(copied_lessons)} lessons copied.")
    return lessons_dir, copied_lessons


def stage_pptx_slides(
    cfg: UnitConfig,
    unit_root: Path,
    lessons_dir: Path,
    logger: logging.Logger,
) -> Tuple[Path, List[Path]]:
    """
    Generate editable PPTX decks directly from lesson JSON, no Canva step.
    Replaces the Canva CSV -> manual-export workflow for units that don't
    want a Canva dependency.
    """
    logger.info("Stage: pptx_slides")

    from cmie.generator.pptx_generator import lesson_json_to_pptx

    slides_dir = unit_root / "slides"
    slides_dir.mkdir(parents=True, exist_ok=True)

    pptx_paths: List[Path] = []
    for lesson_json in sorted(lessons_dir.glob("*.json")):
        out_path = lesson_json_to_pptx(lesson_json, output_dir=slides_dir)
        pptx_paths.append(out_path)

    logger.info(f"  {len(pptx_paths)} PPTX decks generated in {slides_dir}")
    return slides_dir, pptx_paths


def stage_canva_csv(
    cfg: UnitConfig,
    unit_root: Path,
    logger: logging.Logger,
) -> Path:
    logger.info("Stage: canva_csv")

    lesson_root = Path("generated_lessons") / slugify(cfg.title)
    lesson_json_paths = sorted(lesson_root.glob("*/lesson.json"))

    output_dir = unit_root / "04_Slides_CSV"
    output_dir.mkdir(parents=True, exist_ok=True)

    csv_path = output_dir / f"{cfg.unit_id}_canva_slides.csv"

    export_unit_canva_csv(
        unit_title=cfg.title,
        lesson_json_paths=lesson_json_paths,
        output_csv=csv_path,
    )

    logger.info(f"  Canva CSV created: {csv_path}")
    return csv_path


def stage_assessment(
    cfg: UnitConfig,
    unit_root: Path,
    logger: logging.Logger,
) -> Path:
    logger.info("Stage: assessment")

    dest_dir = unit_root / "assessment"
    dest_dir.mkdir(parents=True, exist_ok=True)

    assessment_generator.generate_summative_assessment(
        unit_id=cfg.unit_id,
        title=cfg.title,
        year_level=cfg.year_level,
        subject=cfg.subject,
        output_dir=dest_dir,
    )

    render_assessment_markdown(dest_dir)

    logger.info("  Assessment generated.")
    return dest_dir


def stage_workbook(
    cfg: UnitConfig,
    unit_root: Path,
    lessons_dir: Path,
    assessment_dir: Optional[Path],
    logger: logging.Logger,
) -> Path:
    logger.info("Stage: workbook")

    unit_config_dict = {
        "unit_id": cfg.unit_id,
        "title": cfg.title,
        "year_level": cfg.year_level,
        "subject": cfg.subject,
        "version": cfg.version,
    }

    workbook_path = generate_student_workbook(
        unit_root,
        unit_config_dict,
        lessons_dir,
        assessment_dir,
    )

    logger.info(f"  Workbook file: {workbook_path}")
    return workbook_path

def stage_roadmap(
    cfg: UnitConfig,
    unit_root: Path,
    lessons_dir: Path,
    assessment_dir: Optional[Path],
    logger: logging.Logger,
) -> Path:
    logger.info("Stage: roadmap")

    unit_config_dict = {
        "unit_id": cfg.unit_id,
        "title": cfg.title,
        "year_level": cfg.year_level,
        "subject": cfg.subject,
        "version": cfg.version,
    }

    meta = generate_unit_roadmap(
        unit_root,
        unit_config_dict,
        lessons_dir,
        assessment_dir,
    )
    roadmap_path = Path(meta["file"])

    logger.info(f"  Roadmap file: {roadmap_path}")
    return roadmap_path


def stage_readme(
    cfg: UnitConfig,
    unit_root: Path,
    lessons_dir: Path,
    assessment_dir: Optional[Path],
    logger: logging.Logger,
) -> Path:
    logger.info("Stage: readme")

    unit_config_dict = {
        "unit_id": cfg.unit_id,
        "title": cfg.title,
        "year_level": cfg.year_level,
        "subject": cfg.subject,
        "version": cfg.version,
    }

    readme_path = generate_unit_readme(
        unit_root=unit_root,
        unit_config=unit_config_dict,
        lessons_dir=lessons_dir,
        assessment_dir=assessment_dir,
    )

    logger.info(f"  README file: {readme_path}")
    return readme_path


def stage_marketing(
    cfg: UnitConfig,
    unit_root: Path,
    lessons_dir: Path,
    logger: logging.Logger,
) -> Path:
    logger.info("Stage: marketing")

    dest_dir = unit_root / "marketing"
    dest_dir.mkdir(parents=True, exist_ok=True)

    unit_meta = {
        "unit_id": cfg.unit_id,
        "title": cfg.title,
        "year_level": cfg.year_level,
        "subject": cfg.subject,
        "version": cfg.version,
    }

    assets = generate_marketing_assets(unit_meta, lessons_dir)

    out_path = dest_dir / "marketing_assets.json"
    with out_path.open("w", encoding="utf-8") as f:
        json.dump(assets, f, ensure_ascii=False, indent=2)

    logger.info(f"  Marketing assets generated: {out_path}")
    return out_path


def stage_listings(
    cfg: UnitConfig,
    unit_root: Path,
    logger: logging.Logger,
) -> None:
    logger.info("Stage: listings")

    if generate_listings_for_unit is None:
        logger.info("  Listings generator not available; skipping.")
        return

    unit_config_dict = {
        "unit_id": cfg.unit_id,
        "title": cfg.title,
        "year_level": cfg.year_level,
        "subject": cfg.subject,
        "version": cfg.version,
    }

    try:
        generate_listings_for_unit(unit_root, unit_config_dict)  # type: ignore
        logger.info("  Listings generated.")
    except Exception as e:  # pragma: no cover
        logger.warning(f"  Listings generation failed: {e}")

def render_rubric_table(
    rubric_json_path: Path,
    output_docx: Path,
    logger: logging.Logger,
) -> None:
    """
    Render rubric.json into a proper DOCX table instead of dumping markdown.
    """
    logger.info(f"  Converting {rubric_json_path.name} -> {output_docx.name}")

    with rubric_json_path.open(encoding="utf-8") as f:
        rubric = json.load(f)

    levels = rubric.get("levels", ["Exemplary", "Proficient", "Developing", "Beginning"])
    criteria = rubric.get("criteria", [])

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(11)

    doc.add_paragraph("Assessment rubric and marking guide", style="Title")
    doc.add_paragraph("")
    doc.add_paragraph("Rubric", style="Heading 1")

    table = doc.add_table(rows=len(criteria) + 1, cols=len(levels) + 1)
    table.style = "Table Grid"

    # Header row -- attach a generic A-D grade band to each level so the
    # rubric can actually be used for grading, not just qualitative feedback.
    grade_bands = {"Exemplary": "A", "Proficient": "B", "Developing": "C", "Beginning": "D"}
    table.cell(0, 0).text = "Criterion"
    for i, level in enumerate(levels, start=1):
        level_str = str(level)
        grade = grade_bands.get(level_str)
        table.cell(0, i).text = f"{level_str} ({grade})" if grade else level_str

    # Body rows
    for row_idx, crit in enumerate(criteria, start=1):
        table.cell(row_idx, 0).text = str(crit.get("name", ""))
        descriptors = crit.get("descriptors", {})
        for col_idx, level in enumerate(levels, start=1):
            table.cell(row_idx, col_idx).text = str(descriptors.get(level, ""))

    # Marking guide
    marking_path = rubric_json_path.parent / "marking_guide.json"
    if marking_path.exists():
        with marking_path.open(encoding="utf-8") as f:
            marking = json.load(f)

        doc.add_paragraph("")
        doc.add_paragraph("Marking guide", style="Heading 1")

        general = marking.get("general_notes", [])
        if general:
            doc.add_paragraph("General notes", style="Heading 2")
            for item in general:
                doc.add_paragraph(str(item), style="List Bullet")

        features = marking.get("high_quality_response_features", [])
        if features:
            doc.add_paragraph("Features of high-quality responses", style="Heading 2")
            for item in features:
                doc.add_paragraph(str(item), style="List Bullet")

        misconceptions = marking.get("common_misconceptions", [])
        if misconceptions:
            doc.add_paragraph("Common misconceptions", style="Heading 2")
            for item in misconceptions:
                doc.add_paragraph(str(item), style="List Bullet")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))



def stage_packaging(
    cfg: UnitConfig,
    unit_root: Path,
    releases_root: Path,
    assessment_dir: Path,
    workbook_path: Path,
    roadmap_path: Path,
    readme_path: Path,
    logger: logging.Logger,
) -> None:
    """
    Create:
    - Internal zip of the full unit_root
    - Public editable folder with CSV and DOCX outputs
    """
    logger.info("Stage: packaging")

    # -----------------------------
    # Internal zip
    # -----------------------------
    artifacts_root = releases_root / "artifacts"
    artifacts_root.mkdir(parents=True, exist_ok=True)

    zip_name = f"{cfg.unit_id}_{cfg.version}"
    zip_base = artifacts_root / zip_name
    shutil.make_archive(str(zip_base), "zip", root_dir=unit_root)
    logger.info("  Internal zip created.")


    required_files = [
        assessment_dir / "assessment_task.md",
        assessment_dir / "assessment_rubric_marking.md",
        workbook_path,
        roadmap_path,
        readme_path,
    ]

    missing = [str(p) for p in required_files if not p.exists()]

    if missing:
        raise RuntimeError(f"Missing required files before packaging: {missing}")

    # -----------------------------
    # Public editable folder
    # -----------------------------
    public_root = releases_root / "public" / f"{cfg.unit_id}_{cfg.version}"

    if public_root.exists():
        try:
            shutil.rmtree(public_root)
        except PermissionError as e:
            logger.warning(
                "Public folder is in use (likely a DOCX open in Word). "
                "Skipping public packaging this run.\nDetails: %s",
                e,
            )
            return

    public_root.mkdir(parents=True, exist_ok=True)

    # 01 – Lesson Slides: prefer direct-generated PPTX (no Canva step needed)
    # over the Canva CSV workflow, when PPTX decks exist.
    direct_slides_dir = unit_root / "slides"
    direct_pptx_files = list(direct_slides_dir.glob("*.pptx")) if direct_slides_dir.exists() else []

    if direct_pptx_files:
        public_slides = public_root / "01_Lesson_Slides"
        public_slides.mkdir(parents=True, exist_ok=True)
        for pptx_file in direct_pptx_files:
            shutil.copy2(pptx_file, public_slides / pptx_file.name)
    else:
        public_slides_csv = public_root / "01_Lesson_Slides_CSV"
        public_slides_csv.mkdir(parents=True, exist_ok=True)
        csv_dir = unit_root / "04_Slides_CSV"
        for csv_file in csv_dir.glob("*.csv"):
            shutil.copy2(csv_file, public_slides_csv / csv_file.name)

            # 02 – Assessment
    public_assessment = public_root / "02_Assessment"
    public_assessment.mkdir(parents=True, exist_ok=True)

    task_md = assessment_dir / "assessment_task.md"
    rubric_json = assessment_dir / "rubric.json"

    if task_md.exists():
        markdown_to_docx(
            task_md,
            public_assessment / "Assessment_Task.docx",
            logger,
        )
    else:
        logger.warning(f"Assessment task markdown missing: {task_md}")

    if rubric_json.exists():
        render_rubric_table(
            rubric_json,
            public_assessment / "Assessment_Rubric_and_Marking.docx",
            logger,
        )
    else:
        logger.warning(f"Assessment rubric JSON missing: {rubric_json}")

    # 03 – Student Workbook
    public_workbook = public_root / "03_Student_Workbook"
    public_workbook.mkdir(parents=True, exist_ok=True)

    if workbook_path.exists():
        markdown_to_docx(
            workbook_path,
            public_workbook / "Student_Workbook.docx",
            logger,
        )
    else:
        logger.warning(f"Workbook markdown missing: {workbook_path}")

    # 04 – Unit Roadmap
    public_roadmap = public_root / "04_Unit_Roadmap"
    public_roadmap.mkdir(parents=True, exist_ok=True)

    if roadmap_path.exists():
        markdown_to_docx(
            roadmap_path,
            public_roadmap / "Unit_Roadmap.docx",
            logger,
        )
    else:
        logger.warning(f"Roadmap markdown missing: {roadmap_path}")

    # 05 – Teacher Guide
    public_teacher_guide = public_root / "05_Teacher_Guide"
    public_teacher_guide.mkdir(parents=True, exist_ok=True)

    if readme_path.exists():
        markdown_to_docx(
            readme_path,
            public_teacher_guide / "Teacher_Guide.docx",
            logger,
        )
    else:
        logger.warning(f"Teacher guide markdown missing: {readme_path}")


    # 06 – Listings
    listings_dir = unit_root / "listings"
    if listings_dir.exists():
        public_listings = public_root / "06_Listings"
        public_listings.mkdir(parents=True, exist_ok=True)

        for f in listings_dir.rglob("*.*"):
            if f.is_file():
                dest = public_listings / f.relative_to(listings_dir)
                dest.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(f, dest)


# --------------------------------------------------------------------
# Main
# --------------------------------------------------------------------


def run_pipeline(config_path: Path, releases_root: Path) -> None:
    logger = setup_logger()

    with config_path.open(encoding="utf-8") as f:
        raw_config = json.load(f)

    cfg = UnitConfig.from_dict(raw_config)

    unit_root = releases_root / cfg.unit_id
    unit_root.mkdir(parents=True, exist_ok=True)

    lessons_dir, _ = stage_lessons(cfg, unit_root, logger)
    _slides_dir, _pptx_paths = stage_pptx_slides(cfg, unit_root, lessons_dir, logger)
    assessment_dir = stage_assessment(cfg, unit_root, logger)
    workbook_path = stage_workbook(cfg, unit_root, lessons_dir, assessment_dir, logger)
    roadmap_path = stage_roadmap(cfg, unit_root, lessons_dir, assessment_dir, logger)
    readme_path = stage_readme(cfg, unit_root, lessons_dir, assessment_dir, logger)

    _marketing_path = stage_marketing(cfg, unit_root, lessons_dir, logger)

    try:
        stage_listings(cfg, unit_root, logger)
    except ImportError:
        logger.info("Listings generator not available; skipping.")

    stage_packaging(
        cfg=cfg,
        unit_root=unit_root,
        releases_root=releases_root,
        assessment_dir=assessment_dir,
        workbook_path=workbook_path,
        roadmap_path=roadmap_path,
        readme_path=readme_path,
        logger=logger,
    )

    errors = validate_unit(unit_root)

    report_path = unit_root / "validation_report.md"
    with report_path.open("w", encoding="utf-8") as f:
        f.write(f"# Validation report for {cfg.unit_id} ({cfg.version})\n\n")
        if errors:
            f.write("## Status\n\n")
            f.write("⚠️ Validation completed **with issues**.\n\n")
            f.write("## Issues\n\n")
            for e in errors:
                f.write(f"- {e}\n")
        else:
            f.write("## Status\n\n")
            f.write("✅ Validation completed: no structural issues detected.\n")

    if errors:
        logger.info("Validation completed with issues:")
        for e in errors:
            logger.info(f"  - {e}")
    else:
        logger.info("Validation completed: no structural issues detected.")

    logger.info(f"Validation report written to {report_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Full product pipeline for a micro-unit.")
    parser.add_argument(
        "--unit-config",
        required=True,
        help="Path to the unit config JSON (e.g. data/units/year7_ai_data_unit1.json)",
    )
    parser.add_argument(
        "--releases-root",
        default="releases",
        help="Root folder for releases (default: releases)",
    )
    args = parser.parse_args()

    config_path = Path(args.unit_config)
    releases_root = Path(args.releases_root)

    run_pipeline(config_path, releases_root)


if __name__ == "__main__":
    main()